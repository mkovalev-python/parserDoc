import os
import pathlib
import shutil
import xml.etree.ElementTree
import zipfile

import requests
from docx import Document
import pandas as pd
import time


def get_task(tasks):
    alphabet = 'Й Ц У К Е Н Г Ш Щ З Х Ъ Ф Ы В А П Р О Л Д Ж Э Ё Я Ч С М И Т Ь Б Ю' \
               'й ц у к е н г ш щ з х ъ ф ы в а п р о л д ж э ё я ч с м и т ь б ю'
    list_tasks = []
    for task in tasks:
        i = 0
        while i < task.__len__():
            symbol = task[i:i + 1]
            if symbol in alphabet:
                list_tasks.append(task[i:])
                break
            else:
                i += 1
    task_return = []
    for reverse in list_tasks:
        task_return.append(reverse[::-1])

    return task_return


def clear_text(tasks):
    for task in tasks:
        for text in task['table']:
            if text in task['text']:
                index_start = task['text'].index(text)
                index_stop = task['table'].__len__() + index_start
                del task['text'][index_start:index_stop]
                break
    return tasks


def send_request(user, task, text, file, project):
    """Создаем исполнителей, проект и задачи"""
    response = requests.post('http://127.0.0.1:8000/create-task/',
                             files=file,
                             data={
                                 "user": user,
                                 "task": task,
                                 "text": text,
                                 'project': project
                             })
    if response.status_code == 200:
        print(f'Отправлен файл в задачу {task[:30]}')


def sorted_table(tasks, tables, double_tables, currentFile):
    task_table = []
    for text_table, text_table_d in zip(tables, double_tables):
        for task in tasks:
            if task['text_task'].__len__() == 0:
                continue
            else:
                if text_table[0] in task['text_task']:
                    task_table.append({'task': task['task'], 'table_row': text_table_d, 'table': text_table,
                                       'text': task['text_task']})
                    break

    """очистка текста задач от таблиц"""
    good_text = clear_text(task_table)

    """Создать таблицу и выгрузить в файл"""
    ii = 1
    for table1 in good_text:
        data = pd.DataFrame(table1['table_row'])
        document = Document()

        document.add_heading(table1['task'])
        table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
        for i, column in enumerate(data):
            for row in range(data.shape[0]):
                table.cell(row, i).text = str(data[column][row])
        table.style = 'TableGrid'
        name = table1['task'][:30].replace('Отчет', "") + '.docx'
        try:
            os.mkdir('./tables/' + currentFile.parent.stem.replace('Отчет', ""))
        except FileExistsError:
            pass
        try:
            os.mkdir('./tables/' + currentFile.parent.stem + '/' + currentFile.stem.replace('Отчет', ""))
        except FileExistsError:
            pass
        dir = './tables/' + currentFile.parent.stem + '/' + currentFile.stem.replace('Отчет', "") + '/' + str(ii) + name
        document.save(dir)

        """Из списка в строку с отступами"""
        TEXT = ''
        for text in table1['text']:
            TEXT += f'\n\t{text}'

        file_ob = {'uploaded_file': open(dir, 'rb')}

        send_request(currentFile.parent.stem, table1['task'], TEXT, file_ob, currentFile.stem.replace('Отчет', ""))

        ii += 1


def get_text_in_task(tasks, texts, tables, tables_double, currentFile):
    ij = 1
    list_tasks_texts = []
    for task in tasks:
        for text in texts:
            if task in text:
                i = texts.index(text)
                try:
                    task_2 = tasks[(tasks.index(task) + 1)]
                    for text_2 in texts:
                        if task_2 in text_2:
                            j = texts.index(text_2)
                    text1 = texts[i + 1:j - 1]
                except:
                    text1 = texts[i + 1:]
                TEXT = []

                for text_task in text1:
                    TEXT.append(text_task)

        list_tasks_texts.append({'text_task': TEXT, 'task': task})
    sorted_table(list_tasks_texts, tables, tables_double, currentFile)


def main():
    user_list = ('Бутакова', 'Винтова', 'Грибанова', 'Чернова')
    start_time = time.time()
    for user in user_list:
        currentDirectory = pathlib.Path(f'./files/{user}/')

        for currentFile in currentDirectory.iterdir():
            # todo: Добавить создание проекта

            WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            PARA = WORD_NAMESPACE + 'p'
            TEXT = WORD_NAMESPACE + 't'
            TABLE = WORD_NAMESPACE + 'tbl'
            ROW = WORD_NAMESPACE + 'tr'
            CELL = WORD_NAMESPACE + 'tc'

            list_task = []
            try:
                with zipfile.ZipFile(currentFile) as docx:
                    tree = xml.etree.ElementTree.XML(docx.read('word/document.xml'))
            except zipfile.BadZipFile:
                continue

            """Текст лист"""
            list_text = []
            for text in tree.iter(PARA):
                item = ''.join(node.text for node in text.iter(TEXT))
                if item == '' or item == ' ' or item.__len__() < 2:
                    continue
                else:
                    list_text.append(item)

            for task in list_text:
                if '1.5.Нормативно-правовая база' in task:
                    a = list_text.index(task)
                if 'Настоящий Отчет по результатам анализа принятых регулирующими органами тарифно-балансовых решений за' in task:
                    b = list_text.index(task)

            list_task = list_text[a + 1:b]

            tasks = get_task(list_task)
            """Конечный список задач"""
            tasks = get_task(tasks)
            try:
                c = list_text.index(tasks[0])
            except ValueError:
                for task in list_text:
                    if tasks[0] in task:
                        c = list_text.index(task)
                        break
            texts = list_text[c:]

            """Получение списка таблиц"""
            tables = []
            tables_all_text = []
            i = 0
            for table in tree.iter(TABLE):
                table_text = []
                text_a = []
                for row in table.iter(ROW):
                    text = []
                    for cell in row.iter(CELL):
                        text.append(''.join(node.text for node in cell.iter(TEXT)))
                        text_a.append(''.join(node.text for node in cell.iter(TEXT)))
                    if text.__len__() > 1:
                        table_text.append(text)
                tables_all_text.append(text_a)
                if table_text.__len__() > 0:
                    i += 1
                    name = 'table' + str(i)
                    tables.append(table_text)
            """Получение текста для задачи"""
            get_text_in_task(tasks, texts, tables_all_text, tables, currentFile)
            print(
                f"--- {(time.time() - start_time)} seconds ---\nUSER:{currentFile.parent.name}\nFILE: {currentFile.name}")
        newzipUser = zipfile.ZipFile(
            './tables/' + currentFile.parent.stem + '.zip', 'w')

        newzipUser.write('./tables/' + currentFile.parent.stem)
        newzipUser.close()
        shutil.rmtree('./tables/' + currentFile.parent.stem,
                      ignore_errors=True)


if __name__ == '__main__':
    main()
